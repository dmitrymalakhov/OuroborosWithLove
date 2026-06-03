"""Word/DOCX editing tools.

These tools complement analyze_document: documents.py extracts DOCX text for
understanding, while this module applies user-approved edits to a DOCX copy.
"""

from __future__ import annotations

import datetime
import json
import pathlib
import re
from typing import Any, Dict, Iterable, List, Tuple

from ouroboros.tools.registry import ToolContext, ToolEntry
from ouroboros.utils import safe_relpath

DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
MAX_DOCX_BYTES = 50 * 1024 * 1024
MAX_PREVIEW_PARAGRAPHS = 80
MAX_SEARCH_MATCHES = 120
MIN_EDIT_CONFIDENCE = 0.75


def _require_docx():
    try:
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.text.paragraph import Paragraph

        return Document, OxmlElement, Paragraph
    except Exception as exc:
        raise RuntimeError("python-docx is required for Word editing tools. Install dependencies from requirements.txt.") from exc


def _scope(ctx: ToolContext) -> Dict[str, Any]:
    return {
        "user_id": ctx.current_user_id,
        "user_role": ctx.user_role,
        "drive_root": str(ctx.drive_root),
        "shared_drive_root": str(ctx.shared_drive_root or ctx.drive_root),
    }


def _resolve_workspace_file(ctx: ToolContext, path: str, source: str = "drive") -> Tuple[pathlib.Path, pathlib.Path]:
    if not path or not isinstance(path, str):
        raise ValueError("path must be a non-empty string")

    source = str(source or "drive").strip().lower()
    if source not in {"drive", "repo"}:
        raise ValueError("source must be 'drive' or 'repo'")
    if source == "repo" and str(ctx.user_role or "user").lower() != "admin":
        raise PermissionError("source='repo' is admin-only in multi-user mode")

    root = (ctx.repo_dir if source == "repo" else ctx.drive_root).resolve()
    resolved = (root / safe_relpath(path)).resolve()
    try:
        resolved.relative_to(root)
    except ValueError:
        raise ValueError("Path traversal is not allowed")
    if not resolved.exists():
        raise FileNotFoundError(f"File not found: {path}")
    if not resolved.is_file():
        raise ValueError(f"Path is not a file: {path}")
    return resolved, root


def _resolve_docx_path(ctx: ToolContext, path: str, source: str = "drive") -> Tuple[pathlib.Path, pathlib.Path]:
    resolved, root = _resolve_workspace_file(ctx, path, source)
    if resolved.suffix.lower() != ".docx":
        raise ValueError("Only .docx files are supported by Word editing tools.")
    size = resolved.stat().st_size
    if size > MAX_DOCX_BYTES:
        raise ValueError(f"DOCX is too large: {size} bytes")
    return resolved, root


def _safe_stem(value: str) -> str:
    value = re.sub(r"[^A-Za-z0-9А-Яа-я._ -]+", "-", str(value or "document"))
    value = re.sub(r"\s+", "-", value.strip(" .-_"))
    return value[:80] or "document"


def _safe_output_path(root: pathlib.Path, output_path: str, source_name: str, overwrite: bool) -> Tuple[pathlib.Path, str]:
    if output_path:
        rel = safe_relpath(output_path)
        if pathlib.PurePosixPath(rel).suffix.lower() != ".docx":
            raise ValueError("output_path must end with .docx")
    else:
        stamp = datetime.datetime.now(datetime.timezone.utc).strftime("%Y%m%d-%H%M%S")
        stem = _safe_stem(pathlib.Path(source_name).stem)
        rel = str(pathlib.PurePosixPath("word_edits") / f"{stem}-edited-{stamp}.docx")

    path = (root / rel).resolve()
    try:
        path.relative_to(root.resolve())
    except ValueError:
        raise ValueError("Path traversal is not allowed")
    path.parent.mkdir(parents=True, exist_ok=True)

    if path.exists() and not overwrite:
        stem = path.stem
        suffix = path.suffix
        parent_rel = pathlib.PurePosixPath(rel).parent
        counter = 2
        while path.exists():
            rel = str(parent_rel / f"{stem}-{counter}{suffix}")
            path = (root / rel).resolve()
            counter += 1
    return path, rel


def _clip(text: str, max_chars: int = 500) -> str:
    text = re.sub(r"\s+", " ", str(text or "")).strip()
    if len(text) <= max_chars:
        return text
    return text[:max_chars].rstrip() + "...(truncated)"


def _iter_paragraph_locations(doc: Any) -> Iterable[Dict[str, Any]]:
    for idx, paragraph in enumerate(doc.paragraphs, start=1):
        yield {
            "kind": "paragraph",
            "paragraph": paragraph,
            "location": f"paragraph {idx}",
            "text": paragraph.text or "",
        }

    for table_idx, table in enumerate(doc.tables, start=1):
        for row_idx, row in enumerate(table.rows, start=1):
            for col_idx, cell in enumerate(row.cells, start=1):
                for para_idx, paragraph in enumerate(cell.paragraphs, start=1):
                    yield {
                        "kind": "table_cell",
                        "paragraph": paragraph,
                        "table": table,
                        "row": row,
                        "cell": cell,
                        "location": f"table {table_idx} row {row_idx} col {col_idx} paragraph {para_idx}",
                        "table_index": table_idx,
                        "row_index": row_idx,
                        "col_index": col_idx,
                        "text": paragraph.text or "",
                    }


def _find_text_matches(doc: Any, search_text: str, max_matches: int = MAX_SEARCH_MATCHES) -> List[Dict[str, Any]]:
    if not search_text:
        return []
    matches: List[Dict[str, Any]] = []
    for location in _iter_paragraph_locations(doc):
        text = str(location.get("text") or "")
        start = 0
        occurrence_in_location = 0
        while True:
            idx = text.find(search_text, start)
            if idx < 0:
                break
            occurrence_in_location += 1
            matches.append(
                {
                    **location,
                    "match_index": idx,
                    "occurrence_in_location": occurrence_in_location,
                    "snippet": _clip(text[max(0, idx - 80): idx + len(search_text) + 80], 220),
                }
            )
            if len(matches) >= max_matches:
                return matches
            start = idx + max(1, len(search_text))
    return matches


def _count_matches(text: str, search_text: str) -> int:
    if not search_text:
        return 0
    count = start = 0
    while True:
        idx = text.find(search_text, start)
        if idx < 0:
            return count
        count += 1
        start = idx + max(1, len(search_text))


def _replace_nth(text: str, search_text: str, replacement: str, occurrence: int) -> str:
    count = start = 0
    while True:
        idx = text.find(search_text, start)
        if idx < 0:
            return text
        count += 1
        if count == occurrence:
            return text[:idx] + replacement + text[idx + len(search_text):]
        start = idx + max(1, len(search_text))


def _replace_one_in_runs(paragraph: Any, search_text: str, replacement: str, occurrence: int, allow_reflow: bool) -> str:
    seen = 0
    for run in paragraph.runs:
        run_count = _count_matches(run.text, search_text)
        if seen + run_count >= occurrence:
            run_occurrence = occurrence - seen
            run.text = _replace_nth(run.text, search_text, replacement, run_occurrence)
            return "run"
        seen += run_count

    if search_text in paragraph.text and allow_reflow:
        paragraph.text = _replace_nth(paragraph.text, search_text, replacement, occurrence)
        return "paragraph_reflow"
    if search_text in paragraph.text:
        return "requires_allow_reflow"
    return "not_found"


def _replace_all_in_runs(paragraph: Any, search_text: str, replacement: str, allow_reflow: bool) -> Tuple[int, str]:
    paragraph_count = _count_matches(paragraph.text, search_text)
    if not paragraph_count:
        return 0, "not_found"

    replaced = 0
    for run in paragraph.runs:
        run_count = _count_matches(run.text, search_text)
        if run_count:
            run.text = run.text.replace(search_text, replacement)
            replaced += run_count

    if replaced == paragraph_count:
        return replaced, "run"
    if allow_reflow:
        paragraph.text = paragraph.text.replace(search_text, replacement)
        return paragraph_count, "paragraph_reflow"
    return replaced, "requires_allow_reflow"


def _coerce_operations(operations: Any) -> List[Dict[str, Any]]:
    if isinstance(operations, str):
        operations = json.loads(operations)
    if not isinstance(operations, list):
        raise ValueError("operations must be a list of objects")
    result = []
    for item in operations:
        if not isinstance(item, dict):
            raise ValueError("each operation must be an object")
        result.append(item)
    return result


def _operation_confirmed(op: Dict[str, Any]) -> Tuple[bool, str]:
    if bool(op.get("confirmed")):
        return True, ""
    confidence = op.get("confidence")
    if confidence is None:
        return False, "missing confirmation; pass confirmed=true after user approval"
    try:
        confidence_value = float(confidence)
    except Exception:
        return False, "invalid confidence; pass confirmed=true after user approval"
    if confidence_value < MIN_EDIT_CONFIDENCE:
        return False, f"low confidence {confidence_value:.2f}; ask user or pass confirmed=true"
    return True, ""


def _insert_paragraph_after(paragraph: Any, text: str, style: str = "") -> Any:
    _Document, OxmlElement, Paragraph = _require_docx()
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    new_para.add_run(text)
    return new_para


def _inspect_word_for_edit(
    ctx: ToolContext,
    path: str,
    source: str = "drive",
    search_text: str = "",
    max_matches: int = MAX_SEARCH_MATCHES,
) -> str:
    Document, _OxmlElement, _Paragraph = _require_docx()
    docx_path, _root = _resolve_docx_path(ctx, path, source)
    max_matches = max(1, min(int(max_matches or MAX_SEARCH_MATCHES), MAX_SEARCH_MATCHES))

    ctx.emit_progress_fn(f"Открыл Word-документ `{docx_path.name}`. Анализирую абзацы, таблицы и совпадения.")
    doc = Document(str(docx_path))
    text_locations = list(_iter_paragraph_locations(doc))
    non_empty = [loc for loc in text_locations if str(loc.get("text") or "").strip()]
    matches = _find_text_matches(doc, search_text, max_matches=max_matches) if search_text else []

    preview_lines = [
        f"- {loc['location']}: {loc.get('text') or '[blank]'}"
        for loc in non_empty[:MAX_PREVIEW_PARAGRAPHS]
    ]
    match_lines = [
        f"- {idx}: {match['location']} | snippet={match['snippet']!r}"
        for idx, match in enumerate(matches, start=1)
    ]
    table_lines = [
        f"- table {idx}: rows={len(table.rows)}, cols={len(table.columns)}"
        for idx, table in enumerate(doc.tables, start=1)
    ]

    return "\n".join(
        [
            "# Word Edit Inspection",
            f"- path: {path}",
            f"- file: {docx_path.name}",
            f"- paragraphs: {len(doc.paragraphs)}",
            f"- tables: {len(doc.tables)}",
            "",
            "## Text Preview",
            *(preview_lines or ["- none"]),
            "",
            "## Search Matches",
            *(match_lines or ["- search_text not provided" if not search_text else "- none"]),
            "",
            "## Tables",
            *(table_lines or ["- none"]),
            "",
            "Next: show the user an edit plan with operation/source/confidence before calling edit_word.",
        ]
    )


def _apply_replace_text(doc: Any, op: Dict[str, Any]) -> List[str]:
    search_text = str(op.get("search_text") or "")
    replacement = str(op.get("replacement_text") if op.get("replacement_text") is not None else op.get("text") or "")
    if not search_text:
        return ["replace_text: missing search_text"]

    all_matches = bool(op.get("all_matches"))
    allow_reflow = bool(op.get("allow_reflow"))
    if all_matches:
        if not allow_reflow:
            reflow_required = []
            for location in _iter_paragraph_locations(doc):
                paragraph = location["paragraph"]
                paragraph_count = _count_matches(paragraph.text, search_text)
                run_count = sum(_count_matches(run.text, search_text) for run in paragraph.runs)
                if paragraph_count and paragraph_count != run_count:
                    reflow_required.append(location["location"])
            if reflow_required:
                return [f"replace_text: match spans multiple runs at {', '.join(reflow_required[:5])}; pass allow_reflow=true"]

        applied = []
        for location in _iter_paragraph_locations(doc):
            paragraph = location["paragraph"]
            count, mode = _replace_all_in_runs(paragraph, search_text, replacement, allow_reflow)
            if count and mode != "requires_allow_reflow":
                applied.append(f"replace_text: {count} match(es) in {location['location']} via {mode}")
        if applied:
            return applied
        return [f"replace_text: no match for {search_text!r}"]

    matches = _find_text_matches(doc, search_text, max_matches=MAX_SEARCH_MATCHES)
    if not matches:
        return [f"replace_text: no match for {search_text!r}"]
    try:
        global_occurrence = max(1, int(op.get("occurrence") or 1))
    except Exception:
        global_occurrence = 1
    if global_occurrence > len(matches):
        return [f"replace_text: occurrence {global_occurrence} not found for {search_text!r}"]

    match = matches[global_occurrence - 1]
    mode = _replace_one_in_runs(
        match["paragraph"],
        search_text,
        replacement,
        int(match["occurrence_in_location"]),
        allow_reflow,
    )
    if mode == "requires_allow_reflow":
        return [f"replace_text: match spans multiple runs at {match['location']}; pass allow_reflow=true"]
    if mode == "not_found":
        return [f"replace_text: no match for {search_text!r}"]
    return [f"replace_text: {search_text!r} -> {replacement!r} at {match['location']} via {mode}"]


def _apply_add_paragraph(doc: Any, op: Dict[str, Any]) -> List[str]:
    text = str(op.get("text") or "")
    if not text:
        return ["add_paragraph: missing text"]
    paragraph = doc.add_paragraph(text)
    if op.get("style"):
        paragraph.style = str(op.get("style"))
    return [f"add_paragraph: {text!r}"]


def _apply_add_heading(doc: Any, op: Dict[str, Any]) -> List[str]:
    text = str(op.get("text") or "")
    if not text:
        return ["add_heading: missing text"]
    try:
        level = max(0, min(9, int(op.get("level") or 1)))
    except Exception:
        level = 1
    doc.add_heading(text, level=level)
    return [f"add_heading: level={level} text={text!r}"]


def _apply_insert_after_match(doc: Any, op: Dict[str, Any]) -> List[str]:
    search_text = str(op.get("search_text") or "")
    text = str(op.get("text") or "")
    if not search_text or not text:
        return ["insert_paragraph_after_match: missing search_text or text"]
    matches = _find_text_matches(doc, search_text, max_matches=MAX_SEARCH_MATCHES)
    if not matches:
        return [f"insert_paragraph_after_match: no match for {search_text!r}"]
    try:
        occurrence = max(1, int(op.get("occurrence") or 1))
    except Exception:
        occurrence = 1
    if occurrence > len(matches):
        return [f"insert_paragraph_after_match: occurrence {occurrence} not found for {search_text!r}"]
    match = matches[occurrence - 1]
    _insert_paragraph_after(match["paragraph"], text, str(op.get("style") or ""))
    return [f"insert_paragraph_after_match: inserted after {match['location']}"]


def _apply_set_table_cell(doc: Any, op: Dict[str, Any]) -> List[str]:
    try:
        table_index = max(1, int(op.get("table_index") or 1))
        row_index = max(1, int(op.get("row") or op.get("row_index") or 1))
        col_index = max(1, int(op.get("col") or op.get("col_index") or 1))
    except Exception:
        return ["set_table_cell: table_index, row, and col must be integers"]
    if table_index > len(doc.tables):
        return [f"set_table_cell: table {table_index} not found"]
    table = doc.tables[table_index - 1]
    if row_index > len(table.rows) or col_index > len(table.columns):
        return [f"set_table_cell: cell {row_index},{col_index} not found in table {table_index}"]
    cell = table.cell(row_index - 1, col_index - 1)
    cell.text = str(op.get("text") if op.get("text") is not None else op.get("value") or "")
    return [f"set_table_cell: table {table_index} row {row_index} col {col_index}"]


def _edit_word(
    ctx: ToolContext,
    path: str,
    operations: Any,
    output_path: str = "",
    overwrite: bool = False,
    send_to_chat: bool = True,
) -> str:
    Document, _OxmlElement, _Paragraph = _require_docx()
    docx_path, root = _resolve_docx_path(ctx, path, "drive")
    output, rel_output = _safe_output_path(root, output_path, docx_path.name, bool(overwrite))
    operation_items = _coerce_operations(operations)

    ctx.emit_progress_fn(f"Редактирую копию Word-документа `{docx_path.name}`: {len(operation_items)} операций.")
    doc = Document(str(docx_path))
    applied: List[str] = []
    rejected: List[str] = []

    for idx, op in enumerate(operation_items, start=1):
        op_type = str(op.get("type") or op.get("operation") or "").strip().lower()
        label = str(op.get("field") or op.get("label") or op_type or f"operation_{idx}")
        ok, reason = _operation_confirmed(op)
        if not ok:
            rejected.append(f"{label}: {reason}")
            continue

        try:
            if op_type == "replace_text":
                results = _apply_replace_text(doc, op)
            elif op_type == "add_paragraph":
                results = _apply_add_paragraph(doc, op)
            elif op_type == "add_heading":
                results = _apply_add_heading(doc, op)
            elif op_type == "insert_paragraph_after_match":
                results = _apply_insert_after_match(doc, op)
            elif op_type == "set_table_cell":
                results = _apply_set_table_cell(doc, op)
            else:
                rejected.append(f"{label}: unsupported operation type {op_type!r}")
                continue
        except Exception as exc:
            rejected.append(f"{label}: {type(exc).__name__}: {exc}")
            continue

        real_results = [
            result for result in results
            if "no match" not in result and "not found" not in result and "missing " not in result and "spans multiple runs" not in result
        ]
        if real_results:
            applied.extend(f"{label}: {result}" for result in real_results)
        else:
            rejected.extend(f"{label}: {result}" for result in results)

    if not applied:
        lines = [
            "⚠️ No Word edits were applied",
            f"- source_path: {path}",
            f"- rejected: {len(rejected)}",
            "- output_path: not_created",
            "- telegram_delivery: skipped",
        ]
        if rejected:
            lines.append("## Rejected")
            lines.extend(f"- {item}" for item in rejected[:80])
        return "\n".join(lines)

    doc.save(output)

    queued = False
    if send_to_chat and ctx.current_chat_id:
        ctx.pending_events.append(
            {
                "type": "send_document",
                "chat_id": ctx.current_chat_id,
                "path": rel_output,
                "caption": "Отредактированный Word-документ",
                "filename": output.name,
                "mime_type": DOCX_MIME_TYPE,
                **_scope(ctx),
            }
        )
        queued = True

    lines = [
        "OK: Word document edited",
        f"- source_path: {path}",
        f"- output_path: {rel_output}",
        f"- applied: {len(applied)}",
        f"- rejected: {len(rejected)}",
        f"- telegram_delivery: {'queued' if queued else 'skipped'}",
    ]
    if applied:
        lines.append("## Applied")
        lines.extend(f"- {item}" for item in applied[:80])
    if rejected:
        lines.append("## Rejected")
        lines.extend(f"- {item}" for item in rejected[:80])
    return "\n".join(lines)


def get_tools() -> List[ToolEntry]:
    return [
        ToolEntry(
            "inspect_word_for_edit",
            {
                "name": "inspect_word_for_edit",
                "description": (
                    "Inspect a .docx Word document before editing it. Use this when the user wants corrections "
                    "in a Word file, not merely a summary. It reports text previews, exact search matches, and "
                    "table dimensions. After inspection, show the user an edit plan before calling edit_word."
                ),
                "parameters": {
                    "type": "object",
                    "properties": {
                        "path": {"type": "string", "description": "Path to the DOCX in Drive or repo."},
                        "source": {
                            "type": "string",
                            "enum": ["drive", "repo"],
                            "default": "drive",
                            "description": "Where to inspect the DOCX. repo is admin-only.",
                        },
                        "search_text": {
                            "type": "string",
                            "description": "Optional exact text to locate before planning a replacement or insertion.",
                        },
                        "max_matches": {
                            "type": "integer",
                            "default": MAX_SEARCH_MATCHES,
                            "description": "Maximum search matches to report.",
                        },
                    },
                    "required": ["path"],
                },
            },
            _inspect_word_for_edit,
            timeout_sec=60,
        ),
        ToolEntry(
            "edit_word",
            {
                "name": "edit_word",
                "description": (
                    "Apply confirmed edits to a copy of a .docx Word document and optionally send it to Telegram. "
                    "Supports replace_text, add_paragraph, add_heading, insert_paragraph_after_match, and "
                    "set_table_cell. Inline replacements preserve run formatting when the match is inside one run; "
                    "matches spanning multiple runs require allow_reflow=true. Original DOCX files are never modified."
                ),
                "parameters": {
                    "type": "object",
                    "properties": {
                        "path": {"type": "string", "description": "Drive path to the source DOCX."},
                        "operations": {
                            "type": "array",
                            "description": (
                                "Confirmed edit operations. Each item needs type and confirmed=true or "
                                "confidence>=0.75."
                            ),
                            "items": {
                                "type": "object",
                                "properties": {
                                    "type": {
                                        "type": "string",
                                        "enum": [
                                            "replace_text",
                                            "add_paragraph",
                                            "add_heading",
                                            "insert_paragraph_after_match",
                                            "set_table_cell",
                                        ],
                                    },
                                    "search_text": {"type": "string"},
                                    "replacement_text": {"type": "string"},
                                    "text": {"type": "string"},
                                    "occurrence": {"type": "integer"},
                                    "all_matches": {"type": "boolean"},
                                    "allow_reflow": {"type": "boolean"},
                                    "style": {"type": "string"},
                                    "level": {"type": "integer"},
                                    "table_index": {"type": "integer"},
                                    "row": {"type": "integer"},
                                    "col": {"type": "integer"},
                                    "value": {},
                                    "confidence": {"type": "number"},
                                    "confirmed": {"type": "boolean"},
                                },
                            },
                        },
                        "output_path": {
                            "type": "string",
                            "description": "Optional output .docx path in Drive. Defaults to word_edits/<name>-edited-<timestamp>.docx.",
                        },
                        "overwrite": {
                            "type": "boolean",
                            "default": False,
                            "description": "Overwrite output_path if it exists; otherwise create a deduplicated filename.",
                        },
                        "send_to_chat": {
                            "type": "boolean",
                            "default": True,
                            "description": "Queue the edited DOCX for Telegram delivery when there is an active chat.",
                        },
                    },
                    "required": ["path", "operations"],
                },
            },
            _edit_word,
            timeout_sec=60,
        ),
    ]
