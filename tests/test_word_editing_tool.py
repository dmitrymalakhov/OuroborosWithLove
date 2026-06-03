import pathlib

import pytest
from docx import Document

from ouroboros.tools.registry import ToolContext, ToolRegistry
from ouroboros.tools.word_editing import DOCX_MIME_TYPE, _edit_word, _inspect_word_for_edit


def _ctx(repo_dir: pathlib.Path, drive_root: pathlib.Path) -> ToolContext:
    repo_dir.mkdir(parents=True, exist_ok=True)
    drive_root.mkdir(parents=True, exist_ok=True)
    return ToolContext(
        repo_dir=repo_dir,
        drive_root=drive_root,
        current_chat_id=123,
        current_user_id=456,
        user_role="user",
    )


def _write_docx(path: pathlib.Path) -> None:
    doc = Document()
    doc.add_paragraph("Contract for ACME LLC")
    para = doc.add_paragraph()
    para.add_run("Amount: ")
    amount = para.add_run("100 RUB")
    amount.bold = True
    doc.add_paragraph("Insert after this sentence.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Field"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Manager"
    table.cell(1, 1).text = "Old"
    doc.save(path)


def _paragraph_texts(path: pathlib.Path):
    doc = Document(path)
    return [p.text for p in doc.paragraphs]


def test_inspect_word_for_edit_reports_matches_and_tables(tmp_path):
    drive = tmp_path / "drive"
    repo = tmp_path / "repo"
    drive.mkdir()
    _write_docx(drive / "contract.docx")

    result = _inspect_word_for_edit(_ctx(repo, drive), path="contract.docx", search_text="100 RUB")

    assert "# Word Edit Inspection" in result
    assert "- paragraphs: 3" in result
    assert "- tables: 1" in result
    assert "Contract for ACME LLC" in result
    assert "paragraph 2" in result
    assert "100 RUB" in result
    assert "table 1: rows=2, cols=2" in result


def test_edit_word_replaces_text_preserves_run_formatting_and_queues_delivery(tmp_path):
    drive = tmp_path / "drive"
    repo = tmp_path / "repo"
    drive.mkdir()
    _write_docx(drive / "contract.docx")
    ctx = _ctx(repo, drive)

    result = _edit_word(
        ctx,
        path="contract.docx",
        output_path="exports/contract-edited.docx",
        operations=[
            {
                "type": "replace_text",
                "search_text": "100 RUB",
                "replacement_text": "125 RUB",
                "confirmed": True,
            },
            {
                "type": "insert_paragraph_after_match",
                "search_text": "Insert after this sentence.",
                "text": "Inserted paragraph.",
                "confidence": 0.95,
            },
            {
                "type": "set_table_cell",
                "table_index": 1,
                "row": 2,
                "col": 2,
                "text": "New",
                "confidence": 0.95,
            },
        ],
    )

    assert "OK: Word document edited" in result
    assert "- applied: 3" in result
    assert (drive / "exports" / "contract-edited.docx").exists()
    assert "100 RUB" in "\n".join(_paragraph_texts(drive / "contract.docx"))

    edited = Document(drive / "exports" / "contract-edited.docx")
    assert "125 RUB" in "\n".join(p.text for p in edited.paragraphs)
    assert "Inserted paragraph." in [p.text for p in edited.paragraphs]
    assert edited.tables[0].cell(1, 1).text == "New"
    assert edited.paragraphs[1].runs[1].text == "125 RUB"
    assert edited.paragraphs[1].runs[1].bold is True

    assert len(ctx.pending_events) == 1
    event = ctx.pending_events[0]
    assert event["type"] == "send_document"
    assert event["path"] == "exports/contract-edited.docx"
    assert event["mime_type"] == DOCX_MIME_TYPE


def test_edit_word_does_not_create_output_without_confirmation(tmp_path):
    drive = tmp_path / "drive"
    repo = tmp_path / "repo"
    drive.mkdir()
    _write_docx(drive / "contract.docx")
    ctx = _ctx(repo, drive)

    result = _edit_word(
        ctx,
        path="contract.docx",
        output_path="exports/unconfirmed.docx",
        operations=[
            {"type": "add_paragraph", "text": "Unconfirmed"},
        ],
    )

    assert "No Word edits were applied" in result
    assert "missing confirmation" in result
    assert not (drive / "exports" / "unconfirmed.docx").exists()
    assert ctx.pending_events == []


def test_edit_word_rejects_cross_run_replacement_without_reflow(tmp_path):
    drive = tmp_path / "drive"
    repo = tmp_path / "repo"
    drive.mkdir()
    doc = Document()
    para = doc.add_paragraph()
    para.add_run("Hello ")
    para.add_run("world")
    doc.save(drive / "split.docx")
    ctx = _ctx(repo, drive)

    result = _edit_word(
        ctx,
        path="split.docx",
        output_path="exports/split-edited.docx",
        operations=[
            {
                "type": "replace_text",
                "search_text": "Hello world",
                "replacement_text": "Hi world",
                "confirmed": True,
            }
        ],
    )

    assert "spans multiple runs" in result
    assert not (drive / "exports" / "split-edited.docx").exists()


def test_word_editing_tools_reject_path_traversal(tmp_path):
    drive = tmp_path / "drive"
    repo = tmp_path / "repo"
    drive.mkdir()
    _write_docx(drive / "contract.docx")

    with pytest.raises(ValueError, match="Path traversal"):
        _inspect_word_for_edit(_ctx(repo, drive), path="../contract.docx")


def test_word_editing_pack_is_registered(tmp_path):
    registry = ToolRegistry(repo_dir=tmp_path / "repo", drive_root=tmp_path / "drive")

    tools = registry.get_tools_by_pack("docx", include_dependencies=True)

    assert "inspect_word_for_edit" in tools
    assert "edit_word" in tools
    assert "analyze_document" in tools
